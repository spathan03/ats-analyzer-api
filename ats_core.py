"""
ATS Resume Scoring System - Web Optimized
Modified from original script for web deployment
"""

import os
import sys
import json
import re
from pathlib import Path
from typing import Dict, List, Tuple, Set
from collections import Counter
from datetime import datetime

# Document Processing
try:
    import pdfplumber
    from PyPDF2 import PdfReader
except ImportError:
    pdfplumber = None
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

# NLP Processing
try:
    import spacy
    from spacy.lang.en.stop_words import STOP_WORDS
except ImportError:
    spacy = None
    STOP_WORDS = set()

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
except ImportError:
    TfidfVectorizer = None
    cosine_similarity = None

try:
    from sentence_transformers import SentenceTransformer
    HAS_SENTENCE_TRANSFORMERS = True
except ImportError:
    HAS_SENTENCE_TRANSFORMERS = False
    SentenceTransformer = None

try:
    import language_tool_python
except ImportError:
    language_tool_python = None

# ===========================
# Utility functions
# ===========================

def clean_extracted_text(text: str) -> str:
    """Clean raw text extracted from documents."""
    if not text:
        return ""
    
    # Remove hyphenation at line breaks
    text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)
    text = text.replace('\r', '\n')
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    
    return text.strip()

# ====================================
# DocumentReader (Optimized for web)
# ====================================

class DocumentReader:
    """Handles reading different document formats."""
    
    @staticmethod
    def _read_pdf_with_pdfplumber(filepath: str) -> str:
        """PDF text extraction using pdfplumber."""
        if not pdfplumber:
            return ""
        
        collected = []
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        collected.append(page_text.strip())
        except Exception:
            pass
        return "\n\n".join(collected).strip()

    @staticmethod
    def read_pdf(filepath: str) -> str:
        """Read PDF file."""
        text = DocumentReader._read_pdf_with_pdfplumber(filepath)
        
        # Fallback to PyPDF2 if available
        if (len(text.strip()) < 100 and PdfReader):
            try:
                reader = PdfReader(filepath)
                texts = []
                for page in reader.pages:
                    try:
                        txt = page.extract_text() or ""
                        texts.append(txt)
                    except Exception:
                        continue
                text2 = "\n\n".join(texts).strip()
                if len(text2) > len(text):
                    text = text2
            except Exception:
                pass
        
        return clean_extracted_text(text)

    @staticmethod
    def read_docx(filepath: str) -> str:
        """Read DOCX file."""
        if not Document:
            return ""
        
        try:
            doc = Document(filepath)
            chunks = []
            
            for p in doc.paragraphs:
                if p.text:
                    chunks.append(p.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text.strip())
                    if row_text:
                        chunks.append(" | ".join(row_text))
            
            text = "\n".join(chunks)
            return clean_extracted_text(text)
        except Exception:
            return ""

    @staticmethod
    def read_txt(filepath: str) -> str:
        """Read plain text file."""
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return clean_extracted_text(text)
        except Exception:
            return ""

    @staticmethod
    def read_document(filepath: str) -> str:
        """Auto-detect and read document."""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
        
        ext = Path(filepath).suffix.lower()
        
        if ext == '.pdf':
            return DocumentReader.read_pdf(filepath)
        elif ext == '.docx':
            return DocumentReader.read_docx(filepath)
        elif ext == '.txt':
            return DocumentReader.read_txt(filepath)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

# ===========================
# NLP Processor
# ===========================

class NLPProcessor:
    """Handles NLP operations."""
    
    def __init__(self):
        self.nlp = None
        self.stop_words = STOP_WORDS
        
        if spacy:
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except:
                pass
        
        self.embed_model = None
        if HAS_SENTENCE_TRANSFORMERS and SentenceTransformer:
            try:
                self.embed_model = SentenceTransformer("all-MiniLM-L6-v2")
            except Exception:
                self.embed_model = None
        
        self.technical_skills = {
            'python', 'java', 'javascript', 'c++', 'c#', 'sql', 'aws', 'azure', 'gcp',
            'docker', 'kubernetes', 'react', 'angular', 'vue', 'nodejs', 'django', 'flask',
            'fastapi', 'spring', 'microservices', 'machine learning', 'deep learning', 'ai',
            'data science', 'analytics', 'pandas', 'numpy', 'scikit-learn', 'pytorch',
            'tensorflow', 'spark', 'hadoop', 'kafka', 'airflow', 'tableau', 'power bi',
            'git', 'ci/cd', 'devops', 'rest', 'graphql', 'api', 'linux', 'bash',
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'jira'
        }
        
        self.soft_skills = {
            'leadership', 'communication', 'teamwork', 'problem solving',
            'analytical', 'creative', 'adaptable', 'organized', 'detail-oriented',
            'collaborative', 'strategic', 'innovative', 'motivated', 'ownership',
            'stakeholder management', 'mentoring', 'presentation', 'negotiation'
        }
    
    def extract_keywords(self, text: str, top_n: int = 50) -> List[Tuple[str, float]]:
        """Extract important keywords."""
        if not text.strip() or not self.nlp:
            return []
        
        doc = self.nlp(text.lower())
        phrases = []
        
        for chunk in doc.noun_chunks:
            phrase = chunk.text.strip()
            if 1 <= len(phrase.split()) <= 4 and phrase not in self.stop_words:
                phrases.append(phrase)
        
        for token in doc:
            if (token.pos_ in ['NOUN', 'PROPN', 'ADJ'] and
                token.text.lower() not in self.stop_words and
                len(token.text) > 2 and token.is_alpha):
                phrases.append(token.lemma_.lower())
        
        if not phrases:
            words = [w.lower() for w in re.findall(r'\w+', text) 
                    if w.lower() not in self.stop_words and len(w) > 3]
            phrase_freq = Counter(words)
        else:
            phrase_freq = Counter(phrases)
        
        return phrase_freq.most_common(top_n)
    
    def extract_skills(self, text: str) -> Dict[str, Set[str]]:
        """Extract technical and soft skills."""
        text_lower = text.lower()
        found_technical = {s for s in self.technical_skills if s in text_lower}
        found_soft = {s for s in self.soft_skills if s in text_lower}
        return {
            'technical': found_technical,
            'soft': found_soft
        }
    
    def calculate_similarity(self, text1: str, text2: str) -> float:
        """Semantic similarity between two texts."""
        text1 = text1.strip()
        text2 = text2.strip()
        if not text1 or not text2:
            return 0.0
        
        # Try sentence transformers first
        if self.embed_model is not None:
            try:
                embeddings = self.embed_model.encode([text1, text2], convert_to_tensor=False)
                from numpy import dot
                from numpy.linalg import norm
                v1, v2 = embeddings[0], embeddings[1]
                sim = float(dot(v1, v2) / (norm(v1) * norm(v2) + 1e-9))
                return max(0.0, min(1.0, sim))
            except Exception:
                pass
        
        # Fallback to TF-IDF
        if TfidfVectorizer and cosine_similarity:
            try:
                vectorizer = TfidfVectorizer(stop_words='english', max_features=2000)
                tfidf_matrix = vectorizer.fit_transform([text1, text2])
                similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
                return float(max(0.0, min(1.0, similarity)))
            except Exception:
                pass
        
        return 0.0

# ======================
# Grammar & Readability
# ======================

TECH_TOKENS = {
    'jira', 'bugzilla', 'mysql', 'mssql', 'aws', 'azure', 'git', 'github', 'gitlab',
    'html', 'css', 'javascript', 'python', 'java', 'sql', 'saas'
}

class GrammarChecker:
    """Check grammar and spelling."""
    
    def __init__(self):
        self.tool = None
        if language_tool_python:
            try:
                self.tool = language_tool_python.LanguageTool('en-US')
            except Exception:
                pass
    
    def check(self, text: str) -> List[Dict]:
        """Check for grammar and spelling errors."""
        if not self.tool or not text.strip():
            return []
        
        try:
            matches = self.tool.check(text)
            errors = []
            for match in matches[:50]:
                msg = match.message or ""
                if "British English" in msg:
                    continue
                
                ctx = match.context or ""
                offset = match.offset or 0
                length = match.errorLength or 0
                snippet = ctx[offset:offset + length] if 0 <= offset < len(ctx) else ""
                
                if snippet.isupper() and len(snippet) > 2:
                    continue
                if "@" in snippet or "http" in snippet.lower():
                    continue
                if snippet.lower() in TECH_TOKENS:
                    continue
                
                errors.append({
                    'type': match.ruleId,
                    'message': msg,
                    'context': ctx,
                    'offset': offset,
                    'errorLength': length,
                    'snippet': snippet,
                    'suggestions': match.replacements[:3]
                })
            return errors
        except Exception:
            return []

def estimate_readability(text: str) -> Dict[str, float]:
    """Rough Flesch Reading Ease estimation."""
    if not text.strip():
        return {"score": 0.0, "label": "Unknown"}
    
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    words = re.findall(r'\w+', text)
    
    if not sentences or not words:
        return {"score": 0.0, "label": "Unknown"}
    
    num_sentences = len(sentences)
    num_words = len(words)
    
    vowels = "aeiouy"
    syllables = 0
    for word in words:
        w = word.lower()
        groups = re.findall(r'[aeiouy]+', w)
        syllables += max(1, len(groups))
    
    words_per_sentence = num_words / num_sentences
    syllables_per_word = syllables / num_words
    
    score = 206.835 - 1.015 * words_per_sentence - 84.6 * syllables_per_word
    
    if score >= 70:
        label = "Easy to read"
    elif score >= 50:
        label = "Fairly readable"
    elif score >= 30:
        label = "Difficult"
    else:
        label = "Very difficult"
    
    return {"score": round(score, 1), "label": label}

# ==========================
# Resume Structure Analyzer
# ==========================

class ResumeStructureAnalyzer:
    """Analyze resume structure and formatting."""
    
    REQUIRED_SECTIONS = [
        'experience', 'education', 'skills',
        'work experience', 'employment', 'professional experience'
    ]
    
    CONTACT_PATTERNS = [
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        r'\b(?:\+?\d{1,3}[-.\s]?)?(?:\d{3}[-.\s]?\d{3}[-.\s]?\d{4})\b',
        r'\blinkedin\.com/in/[A-Za-z0-9\-_/]+\b'
    ]
    
    BULLET_PREFIXES = ('•', '-', '*', '·', '○', '▪', '➢', '▶', '→', '■', '‣')
    
    def analyze(self, text: str) -> Dict:
        """Analyze resume structure."""
        text_lower = text.lower()
        lines = [l for l in text.split('\n') if l.strip()]
        
        sections_found = []
        for section in self.REQUIRED_SECTIONS:
            if section in text_lower:
                sections_found.append(section)
        
        has_email = bool(re.search(self.CONTACT_PATTERNS[0], text))
        has_phone = bool(re.search(self.CONTACT_PATTERNS[1], text))
        has_linkedin = bool(re.search(self.CONTACT_PATTERNS[2], text.lower()))
        
        bullet_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped.startswith(self.BULLET_PREFIXES):
                bullet_lines.append(line)
        
        date_pattern = r'\b(19|20)\d{2}\b|\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b'
        dates_found = len(re.findall(date_pattern, text, re.IGNORECASE))
        
        return {
            'sections_found': list(set(sections_found)),
            'sections_missing': [s for s in ['experience', 'education', 'skills'] 
                                if not any(s in sf for sf in sections_found)],
            'has_contact': {
                'email': has_email,
                'phone': has_phone,
                'linkedin': has_linkedin
            },
            'has_bullets': len(bullet_lines) > 0,
            'bullet_count': len(bullet_lines),
            'dates_found': dates_found,
            'has_chronology': dates_found >= 2,
            'line_count': len(lines)
        }

# ================
# ATS Scoring Core
# ================

class ATSScorer:
    """Main ATS scoring engine."""
    
    def __init__(self):
        self.nlp_processor = NLPProcessor()
        self.grammar_checker = GrammarChecker()
        self.structure_analyzer = ResumeStructureAnalyzer()
    
    def score_keyword_matching(self, resume_text: str, jd_text: str) -> Dict:
        """Score: 30% - Keyword matching."""
        jd_keywords_list = self.nlp_processor.extract_keywords(jd_text, 40)
        resume_keywords_list = self.nlp_processor.extract_keywords(resume_text, 80)
        
        jd_keywords = {k for k, _ in jd_keywords_list}
        resume_keywords = {k for k, _ in resume_keywords_list}
        
        if not jd_keywords:
            return {
                'score': 1.5,
                'max_score': 3.0,
                'percentage': 0.0,
                'matched': [],
                'missing': [],
                'total_jd_keywords': 0,
                'total_matched': 0
            }
        
        matched_keywords = sorted(list(jd_keywords.intersection(resume_keywords)))
        missing_keywords = sorted(list(jd_keywords.difference(resume_keywords)))
        
        match_ratio = len(matched_keywords) / len(jd_keywords)
        score = match_ratio * 3.0
        
        return {
            'score': round(score, 2),
            'max_score': 3.0,
            'percentage': round(match_ratio * 100, 1),
            'matched': matched_keywords[:15],
            'missing': missing_keywords[:15],
            'total_jd_keywords': len(jd_keywords),
            'total_matched': len(matched_keywords)
        }
    
    def score_skills_match(self, resume_text: str, jd_text: str) -> Dict:
        """Score: 20% - Skills matching."""
        resume_skills = self.nlp_processor.extract_skills(resume_text)
        jd_skills = self.nlp_processor.extract_skills(jd_text)
        
        matched_technical = resume_skills['technical'].intersection(jd_skills['technical'])
        missing_technical = jd_skills['technical'] - resume_skills['technical']
        
        matched_soft = resume_skills['soft'].intersection(jd_skills['soft'])
        missing_soft = jd_skills['soft'] - resume_skills['soft']
        
        total_jd_skills = len(jd_skills['technical']) + len(jd_skills['soft'])
        total_matched = len(matched_technical) + len(matched_soft)
        
        match_ratio = total_matched / total_jd_skills if total_jd_skills > 0 else 0.0
        score = match_ratio * 2.0
        
        return {
            'score': round(score, 2),
            'max_score': 2.0,
            'percentage': round(match_ratio * 100, 1),
            'matched_technical': sorted(list(matched_technical)),
            'missing_technical': sorted(list(missing_technical)),
            'matched_soft': sorted(list(matched_soft)),
            'missing_soft': sorted(list(missing_soft))
        }
    
    def score_structure(self, resume_text: str) -> Dict:
        """Score: 15% - Resume structure."""
        structure = self.structure_analyzer.analyze(resume_text)
        
        score = 0.0
        max_score = 1.5
        feedback = []
        
        sections_score = len(structure['sections_found']) / 3 * 0.5
        score += min(0.5, sections_score)
        if structure['sections_missing']:
            feedback.append(f"Missing sections: {', '.join(structure['sections_missing'])}")
        
        contact_count = sum(structure['has_contact'].values())
        contact_score = contact_count / 3 * 0.3
        score += min(0.3, contact_score)
        if contact_count < 3:
            missing = [k for k, v in structure['has_contact'].items() if not v]
            if missing:
                feedback.append(f"Missing contact info: {', '.join(missing)}")
        
        bullet_score = 0.3 if structure['has_bullets'] else 0.0
        score += bullet_score
        if not structure['has_bullets']:
            feedback.append("No bullet points found - use bullets for responsibilities and achievements.")
        
        chrono_score = 0.4 if structure['has_chronology'] else 0.0
        score += chrono_score
        if not structure['has_chronology']:
            feedback.append("Add dates to work experience and education for clear chronology.")
        
        return {
            'score': round(score, 2),
            'max_score': max_score,
            'percentage': round((score / max_score) * 100, 1) if max_score > 0 else 0.0,
            'structure': structure,
            'feedback': feedback
        }
    
    def score_grammar(self, resume_text: str) -> Dict:
        """Score: 15% - Grammar and readability."""
        errors = self.grammar_checker.check(resume_text)
        error_count = len(errors)
        
        if error_count == 0:
            score = 1.5
        elif error_count <= 5:
            score = 1.2
        elif error_count <= 10:
            score = 0.9
        elif error_count <= 15:
            score = 0.6
        else:
            score = 0.3
        
        readability = estimate_readability(resume_text)
        
        return {
            'score': round(score, 2),
            'max_score': 1.5,
            'error_count': error_count,
            'errors': errors[:10],
            'feedback': f"Found {error_count} grammar/spelling issues" if error_count > 0 else "Good grammar",
            'readability_score': readability['score'],
            'readability_label': readability['label']
        }
    
    def score_experience_relevance(self, resume_text: str, jd_text: str) -> Dict:
        """Score: 10% - Experience relevance."""
        similarity = self.nlp_processor.calculate_similarity(resume_text, jd_text)
        score = similarity * 1.0
        
        feedback = []
        if similarity < 0.3:
            feedback.append("Resume content doesn't closely match the job requirements.")
        elif similarity < 0.6:
            feedback.append("Resume has moderate relevance to the job description.")
        else:
            feedback.append("Resume content closely matches the job requirements.")
        
        return {
            'score': round(score, 2),
            'max_score': 1.0,
            'similarity': round(similarity, 3),
            'percentage': round(similarity * 100, 1),
            'feedback': feedback
        }
    
    def score_missing_requirements(self, resume_text: str, jd_text: str) -> Dict:
        """Score: 10% - Missing critical requirements."""
        jd_lower = jd_text.lower()
        resume_lower = resume_text.lower()
        
        required_pattern = r'(?:required|must have|mandatory|must-have)[:\s]+([^\n.]+)'
        required_matches = re.findall(required_pattern, jd_lower)
        
        if not required_matches:
            return {
                'score': 0.8,
                'max_score': 1.0,
                'missing_count': 0,
                'critical_missing': [],
                'feedback': "No explicit 'must have' requirements detected in JD."
            }
        
        critical_missing = []
        for req in required_matches:
            words = [w for w in re.findall(r'\w+', req) if len(w) > 3]
            found = any(word in resume_lower for word in words[:4])
            if not found:
                critical_missing.append(req.strip()[:80])
        
        missing_count = len(critical_missing)
        if missing_count == 0:
            score = 1.0
        elif missing_count <= 2:
            score = 0.7
        elif missing_count <= 4:
            score = 0.4
        else:
            score = 0.1
        
        return {
            'score': round(score, 2),
            'max_score': 1.0,
            'missing_count': missing_count,
            'critical_missing': critical_missing[:5],
            'feedback': f"Missing {missing_count} critical requirements" if missing_count > 0 else "All critical requirements appear to be covered."
        }
    
    def generate_recommendations(self, analysis: Dict) -> List[str]:
        """Generate actionable recommendations."""
        recommendations = []
        
        if analysis['keyword_matching']['total_jd_keywords'] > 0 and analysis['keyword_matching']['percentage'] < 70:
            missing = analysis['keyword_matching']['missing'][:7]
            if missing:
                recommendations.append(f"📌 Add or naturally incorporate these job-related keywords (if relevant): {', '.join(missing)}")
        
        skills = analysis['skills_match']
        if skills['missing_technical']:
            recommendations.append(f"🔧 Highlight or add technical skills (if you actually have them): {', '.join(list(skills['missing_technical'])[:5])}")
        if skills['missing_soft']:
            recommendations.append(f"💡 Emphasize soft skills like: {', '.join(list(skills['missing_soft'])[:3])}")
        
        for fb in analysis['structure']['feedback'][:4]:
            recommendations.append(f"📋 {fb}")
        
        if analysis['grammar']['error_count'] > 0:
            recommendations.append(f"✏️ Fix {analysis['grammar']['error_count']} grammar/spelling issues.")
        if analysis['grammar']['readability_score'] < 50:
            recommendations.append(f"📖 Improve readability ({analysis['grammar']['readability_label']}).")
        
        if analysis['experience_relevance']['similarity'] < 0.5:
            recommendations.append("📝 Tailor your experience bullets to mirror responsibilities and skills in the job description.")
        
        if analysis['missing_requirements']['missing_count'] > 0:
            recommendations.append("⚠️ Explicitly address the 'required' or 'must have' items where you genuinely meet them.")
        
        return recommendations
    
    def calculate_total_score(self, analysis: Dict) -> float:
        """Calculate total ATS score out of 10."""
        total = 0.0
        total += analysis['keyword_matching']['score']
        total += analysis['skills_match']['score']
        total += analysis['structure']['score']
        total += analysis['grammar']['score']
        total += analysis['experience_relevance']['score']
        total += analysis['missing_requirements']['score']
        return round(total, 2)
    
    def analyze_resume(self, resume_text: str, jd_text: str) -> Dict:
        """Complete resume analysis."""
        analysis = {
            'keyword_matching': self.score_keyword_matching(resume_text, jd_text),
            'skills_match': self.score_skills_match(resume_text, jd_text),
            'structure': self.score_structure(resume_text),
            'grammar': self.score_grammar(resume_text),
            'experience_relevance': self.score_experience_relevance(resume_text, jd_text),
            'missing_requirements': self.score_missing_requirements(resume_text, jd_text)
        }
        
        analysis['total_score'] = self.calculate_total_score(analysis)
        analysis['recommendations'] = self.generate_recommendations(analysis)
        
        return analysis